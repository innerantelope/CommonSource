import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding (in dxa: 20 dxa = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_page_number(run):
    """Add page number field to a run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def create_document():
    doc = docx.Document()

    # Page setup - Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Configure footer
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run("CommonSource Overview | Page ")
        f_run.font.name = 'Arial'
        f_run.font.size = Pt(9)
        f_run.font.italic = True
        f_run.font.color.rgb = RGBColor(128, 128, 128)
        add_page_number(f_p.add_run())

    # Styling colors
    PRIMARY_COLOR = RGBColor(26, 54, 93)   # Deep Navy
    SECONDARY_COLOR = RGBColor(43, 108, 176) # Steel Blue
    TEXT_COLOR = RGBColor(45, 55, 72)      # Charcoal
    LIGHT_GRAY = "F7FAFC"
    DARK_BLUE_HEX = "1A365D"
    STEEL_BLUE_HEX = "2B6CB0"
    BORDER_GRAY_HEX = "E2E8F0"

    # Set base font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = TEXT_COLOR

    # Title Page / Header Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("CommonSource (Project_D)")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = PRIMARY_COLOR

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(24)
    sub_run = sub_p.add_run("AI-Powered Community Media Archive Search & Evidence Retrieval System")
    sub_run.font.size = Pt(14)
    sub_run.font.italic = True
    sub_run.font.color.rgb = SECONDARY_COLOR

    # Add a divider line
    p_div = doc.add_paragraph()
    p_div_run = p_div.add_run("―" * 45)
    p_div_run.font.color.rgb = SECONDARY_COLOR
    p_div.paragraph_format.space_after = Pt(24)

    # 1. Executive Summary
    h1 = doc.add_heading(level=1)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    run = h1.add_run("1. Executive Summary")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    p = doc.add_paragraph(
        "CommonSource is a community-focused, AI-powered media archive search and evidence retrieval system designed to "
        "empower local organizations, journalists, and non-governmental organizations (NGOs). By indexing local news feeds, "
        "PDF collections, audio transcriptions, and word documents, CommonSource preserves community institutional knowledge "
        "and makes it searchable via vector semantic matching and traditional lexical search."
    )
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15

    p2 = doc.add_paragraph(
        "Unlike generic search engines or consumer LLMs that suffer from hallucination and lack local context, CommonSource "
        "combines local search databases (SQLite and Qdrant) with retrieval-augmented generation (RAG) powered by Gemini and local "
        "models (like Qwen via Ollama). This allows users to generate grounded syntheses, track narrative storylines over time, "
        "and interactively draft radio or podcast scripts backed by cited community journalism."
    )
    p2.paragraph_format.space_after = Pt(18)
    p2.paragraph_format.line_spacing = 1.15

    # 2. Key Features
    h2 = doc.add_heading(level=1)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    run = h2.add_run("2. Key Product Capabilities")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    features = [
        ("Hybrid Vector & Lexical Search", "Integrates semantic vector matching (MiniLM embeddings via Qdrant/SQLite) with BM25-like lexical keyword scoring, returning relevant documentation even when search terms do not match exactly."),
        ("PageRank Citation Graph Boosting", "Analyzes reference patterns and citations between articles to construct a citation network using NetworkX. Highly cited documents receive up to a 10% search rank boost, surfacing the most authoritative community sources first."),
        ("Multi-Perspective Evidence Layering", "Categorizes search results into five distinct evidence layers (News, Governance/Development, Community, Academic, and Official/PR). The system synthesizes these angles using LLMs to present diverse viewpoints and highlight news coverage gaps."),
        ("Interactive Script Writer", "An LLM-driven collaborative assistant that reads top retrieved documents and allows users to draft scripts (for radio, podcasts, or newsletters) while preserving strict grounding and citations to source material."),
        ("Chronological Timeline & Story Arcs", "Analyzes publication dates and topics to display histogram-based timelines of narratives and automatically extract narrative progression ('Story Arcs') over several years."),
        ("Localized Translation (English & Hindi)", "Provides direct English-to-Hindi and Hindi-to-English translation of search results and synthesized answers, ensuring access for community leaders who communicate in regional languages."),
        ("Secure Authentication & Role Workflows", "Features a comprehensive role-based access control system (Super Admin, Admin, Publisher, Reviewer, Reader) with JWT token-pairs, CSRF protection, and lockout policies.")
    ]

    for title, desc in features:
        p_feat = doc.add_paragraph(style='List Bullet')
        p_feat.paragraph_format.space_after = Pt(6)
        p_feat.paragraph_format.line_spacing = 1.15
        r_bold = p_feat.add_run(f"{title}: ")
        r_bold.font.bold = True
        r_bold.font.color.rgb = SECONDARY_COLOR
        p_feat.add_run(desc)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(12)

    # 3. System Architecture & Tech Stack
    h3 = doc.add_heading(level=1)
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)
    run = h3.add_run("3. System Architecture & Tech Stack")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    p_arch = doc.add_paragraph(
        "CommonSource is structured as a modular Flask application backed by durable storage engines. It consists of a decoupled "
        "frontend, a semantic retrieval service, and a metadata hydration layer. Below is the list of technical components in the stack:"
    )
    p_arch.paragraph_format.space_after = Pt(12)

    # Table of Tech Stack
    tech_table = doc.add_table(rows=6, cols=3)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Headers
    headers = ["Layer", "Technologies Used", "Functionality"]
    for i, h_text in enumerate(headers):
        cell = tech_table.cell(0, i)
        cell.text = h_text
        set_cell_background(cell, DARK_BLUE_HEX)
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        p_cell = cell.paragraphs[0]
        p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p_cell.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(10)

    tech_data = [
        ("Frontend", "HTML5, Vanilla CSS, Vanilla JavaScript", "Provides search UI, interactive script-writing console, timeline charts, and translation toggles without external build steps or node dependencies."),
        ("Application Server", "Python, Flask 3.x, Flask-CORS", "Exposes search, translation, script generation, and authentication APIs. Manages application state and routes requests to backends."),
        ("Relational Storage", "SQLite 3", "Houses the primary knowledge asset database, user profiles, article metadata, offline world model extractions, and precomputed PageRank scores."),
        ("Vector DB & Models", "Qdrant Vector DB, Sentence-Transformers (MiniLM)", "Generates 384-dimensional text embeddings locally and performs fast approximate nearest neighbors (ANN) vector search."),
        ("LLM Reasoning / RAG", "Gemini 2.5 Flash, Ollama (Qwen2.5 / Gemma3)", "Performs text generation, RAG document synthesis, script authoring, entity translation, and multi-perspective gap analysis.")
    ]

    for row_idx, (layer, tech, func) in enumerate(tech_data, start=1):
        for col_idx, text in enumerate([layer, tech, func]):
            cell = tech_table.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if row_idx % 2 == 0:
                set_cell_background(cell, "F7FAFC")
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.line_spacing = 1.15
            for r in p_cell.runs:
                r.font.size = Pt(9.5)
                if col_idx == 0:
                    r.font.bold = True

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_after = Pt(18)

    # 4. Storage & DB Schema
    h4 = doc.add_heading(level=1)
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(6)
    run = h4.add_run("4. Relational Database Schema & Data Models")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    p_db = doc.add_paragraph(
        "CommonSource uses SQLite to maintain rich structural links and metadata. Important tables in the commonsource.db include:"
    )
    p_db.paragraph_format.space_after = Pt(12)

    db_tables = [
        ("knowledge_assets", "The parent table representing any raw document, file, or feed article indexed into the system."),
        ("commonsource_articles", "Contains metadata and provenance of journalism articles, including publisher information, authorship, publish dates, family, content types, and evidence layers."),
        ("knowledge_chunks", "Stores indexed text chunks alongside their binary vector embeddings (used for legacy SQLite-based vector fallback)."),
        ("pagerank_scores", "Precomputed PageRank values of sources computed from citation linkages, updated offline to avoid runtime delays."),
        ("domain_classifications", "Maps documents to domain-specific taxonomy packs (e.g. water, health, climate) to filter searches."),
        ("knowledge_extractions", "Stores JSON structures representing LLM-extracted actors, constraints, stressors, thresholds, and dilemmas from documents."),
        ("causal_network & approved_world_models", "Maintains linkages of systems-dynamics logic (e.g. A causes B) extracted from historical community coverage to model how issues interact.")
    ]

    for table_name, desc in db_tables:
        p_t = doc.add_paragraph(style='List Bullet')
        p_t.paragraph_format.space_after = Pt(6)
        p_t.paragraph_format.line_spacing = 1.15
        r_bold = p_t.add_run(f"{table_name}: ")
        r_bold.font.bold = True
        r_bold.font.color.rgb = SECONDARY_COLOR
        p_t.add_run(desc)

    p_space3 = doc.add_paragraph()
    p_space3.paragraph_format.space_after = Pt(12)

    # 5. Ingestion Pipeline
    h5 = doc.add_heading(level=1)
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(6)
    run = h5.add_run("5. Ingestion & Search Retrieval Flow")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    # Ingestion Workflow
    h5_sub1 = doc.add_heading(level=2)
    h5_sub1.paragraph_format.space_before = Pt(12)
    h5_sub1.paragraph_format.space_after = Pt(4)
    run_s1 = h5_sub1.add_run("5.1 Ingestion Flow")
    run_s1.font.size = Pt(13)
    run_s1.font.bold = True
    run_s1.font.color.rgb = SECONDARY_COLOR

    p_ing = doc.add_paragraph(
        "CommonSource processes files and RSS feeds offline or via ingestion endpoints. The pipeline operates as follows:\n"
        "1. Extraction: Document text is parsed from PDF, DOCX, audio transcripts, or RSS feeds.\n"
        "2. Chunking & Classification: Text is broken into coherent chunks. The text source is classified using source_classifier.py into evidence layers.\n"
        "3. Embedding: Text chunks are converted into 384-dimensional dense vectors using a local sentence-transformer model (all-MiniLM-L6-v2) and cached to disk.\n"
        "4. Database Entry: Chunks are stored in SQLite; vector indices and metadata are synchronized into Qdrant collection 'commonsource_chunks'.\n"
        "5. Graph Analysis: PageRank scripts compute the citations graph from text patterns and write scores into SQLite."
    )
    p_ing.paragraph_format.space_after = Pt(12)
    p_ing.paragraph_format.line_spacing = 1.15

    # Retrieval Workflow
    h5_sub2 = doc.add_heading(level=2)
    h5_sub2.paragraph_format.space_before = Pt(12)
    h5_sub2.paragraph_format.space_after = Pt(4)
    run_s2 = h5_sub2.add_run("5.2 Retrieval Flow")
    run_s2.font.size = Pt(13)
    run_s2.font.bold = True
    run_s2.font.color.rgb = SECONDARY_COLOR

    p_ret = doc.add_paragraph(
        "When a user performs a search or asks a question:\n"
        "1. Query Embedding: The system embedding service (embed.py) converts the query into a 384-dimensional vector.\n"
        "2. Vector Match: Qdrant performs an Approximate Nearest Neighbor (ANN) search to find top matching text chunks. If Qdrant is offline, the system falls back to a SQLite linear search over pre-filtered candidate rows.\n"
        "3. Hydration & Hybrid Scoring: SQLite hydrates document metadata. Scoring.py combines semantic similarity (45%), keyword lexical relevance (45%), and PageRank authority boost (up to 10%).\n"
        "4. RAG Synthesizer: Top ranked chunks are compiled into a prompt context, and sent to Gemini / Ollama Qwen to synthesize a comprehensive reply with inline citations."
    )
    p_ret.paragraph_format.space_after = Pt(18)
    p_ret.paragraph_format.line_spacing = 1.15

    # 6. Setup & Verification
    h6 = doc.add_heading(level=1)
    h6.paragraph_format.space_before = Pt(18)
    h6.paragraph_format.space_after = Pt(6)
    run = h6.add_run("6. Setup, Configuration & Verification")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    p_setup = doc.add_paragraph(
        "To run CommonSource locally, configure the environment variables in a .env file and run the launch commands. "
        "The following parameters are supported:"
    )
    p_setup.paragraph_format.space_after = Pt(12)

    # Config parameters Table
    config_table = doc.add_table(rows=6, cols=3)
    config_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers_conf = ["Environment Variable", "Default Value", "Description"]
    for i, h_text in enumerate(headers_conf):
        cell = config_table.cell(0, i)
        cell.text = h_text
        set_cell_background(cell, STEEL_BLUE_HEX)
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        p_cell = cell.paragraphs[0]
        p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p_cell.runs:
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(10)

    config_data = [
        ("COMMONSOURCE_PORT", "5050", "Port on which the backend Flask app runs."),
        ("COMMONSOURCE_USE_QDRANT", "false", "Set to true to use Qdrant for semantic search. Otherwise falls back to SQLite."),
        ("QDRANT_URL", "http://localhost:6333", "Connection string for Qdrant database server."),
        ("COMMONSOURCE_LLM_PROVIDER", "ollama", "LLM reasoning provider. Options: ollama, gemini, groq, openrouter, auto."),
        ("COMMONSOURCE_LLM_MODEL", "gemini-2.5-flash", "Target LLM model used for synthesis, scripts, and translation.")
    ]

    for row_idx, (env_var, default_v, desc_v) in enumerate(config_data, start=1):
        for col_idx, text in enumerate([env_var, default_v, desc_v]):
            cell = config_table.cell(row_idx, col_idx)
            cell.text = text
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            if row_idx % 2 == 0:
                set_cell_background(cell, "F7FAFC")
            p_cell = cell.paragraphs[0]
            p_cell.paragraph_format.line_spacing = 1.15
            for r in p_cell.runs:
                r.font.size = Pt(9.5)
                if col_idx == 0:
                    r.font.bold = True

    p_space4 = doc.add_paragraph()
    p_space4.paragraph_format.space_after = Pt(12)

    p_cmd = doc.add_paragraph(
        "Important commands to run the system components:\n"
        "• Start Flask server: python Project/app/search_api.py\n"
        "• Sync Qdrant vectors: python Project/scripts/sync_qdrant.py --recreate\n"
        "• Compute PageRank: python Project/scripts/compute_pagerank.py\n"
        "• Expose demo with tunnel: cloudflared tunnel --url http://localhost:5050"
    )
    p_cmd.paragraph_format.space_after = Pt(18)
    p_cmd.paragraph_format.line_spacing = 1.15

    # Footer/Sign off
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.space_before = Pt(36)
    r_foot = p_footer.add_run("CommonSource Technical Specification & Architecture Overview Document")
    r_foot.font.italic = True
    r_foot.font.size = Pt(9.5)
    r_foot.font.color.rgb = RGBColor(160, 174, 192)

    doc.save("CommonSource_Overview.docx")
    print("Document saved successfully as CommonSource_Overview.docx")

if __name__ == "__main__":
    create_document()
